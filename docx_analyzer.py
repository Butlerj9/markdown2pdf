#!/usr/bin/env python3
"""
DOCX Analyzer for Markdown to PDF Converter Verification
-------------------------------------------------------
This script analyzes DOCX files to extract settings information for verification.
"""

import os
import re
import logging
from docx import Document
from docx.shared import Pt, Inches

# Configure logger
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger('docx_analyzer')

def analyze_docx(docx_file):
    """
    Analyze a DOCX file to extract settings information
    
    Args:
        docx_file (str): Path to the DOCX file
        
    Returns:
        dict: Dictionary containing extracted settings
    """
    logger.info(f"Analyzing DOCX file: {docx_file}")
    
    if not os.path.exists(docx_file):
        logger.error(f"DOCX file not found: {docx_file}")
        return {"error": f"DOCX file not found: {docx_file}"}
    
    try:
        # Open the document
        doc = Document(docx_file)
        
        # Extract document properties
        doc_properties = extract_document_properties(doc)
        
        # Extract page setup
        page_setup = extract_page_setup(doc)
        
        # Extract font settings
        font_info = extract_font_settings(doc)
        
        # Check for TOC presence
        toc_info = check_toc_presence(doc)
        
        # Check heading numbering
        heading_info = check_heading_numbering(doc)
        
        # Combine all information
        result = {
            **doc_properties,
            **page_setup,
            **font_info,
            **toc_info,
            **heading_info
        }
        
        logger.info(f"DOCX analysis complete: {docx_file}")
        return result
    
    except Exception as e:
        logger.error(f"Error analyzing DOCX file: {str(e)}")
        return {"error": f"Error analyzing DOCX file: {str(e)}"}

def extract_document_properties(doc):
    """Extract document properties"""
    try:
        properties = {
            "title": doc.core_properties.title,
            "author": doc.core_properties.author,
            "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None,
            "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None
        }
        
        return {
            "document_properties": properties
        }
    
    except Exception as e:
        logger.error(f"Error extracting document properties: {str(e)}")
        return {"error": f"Error extracting document properties: {str(e)}"}

def extract_page_setup(doc):
    """Extract page setup information"""
    try:
        # Get page setup from first section
        section = doc.sections[0]
        
        # Get page dimensions
        width = section.page_width.inches
        height = section.page_height.inches
        
        # Determine page size and orientation
        page_size, orientation = determine_page_size(width, height)
        
        # Get margins
        margins = {
            "top": section.top_margin.inches,
            "right": section.right_margin.inches,
            "bottom": section.bottom_margin.inches,
            "left": section.left_margin.inches
        }
        
        return {
            "page_size": page_size,
            "orientation": orientation,
            "width_inches": width,
            "height_inches": height,
            "margins": margins
        }
    
    except Exception as e:
        logger.error(f"Error extracting page setup: {str(e)}")
        return {"error": f"Error extracting page setup: {str(e)}"}

def determine_page_size(width, height):
    """
    Determine page size and orientation based on dimensions
    
    Common page sizes in inches:
    - A4: 8.27 x 11.69 (portrait), 11.69 x 8.27 (landscape)
    - Letter: 8.5 x 11 (portrait), 11 x 8.5 (landscape)
    """
    # Normalize by making width the smaller dimension
    if width > height:
        w, h = height, width
        orientation = "landscape"
    else:
        w, h = width, height
        orientation = "portrait"
    
    # Check common page sizes with some tolerance
    tolerance = 0.1
    
    # A4
    if abs(w - 8.27) <= tolerance and abs(h - 11.69) <= tolerance:
        return "A4", orientation
    
    # Letter
    if abs(w - 8.5) <= tolerance and abs(h - 11.0) <= tolerance:
        return "Letter", orientation
    
    # If no match, return custom
    return f"Custom ({width:.2f}x{height:.2f})", orientation

def extract_font_settings(doc):
    """Extract font settings from document"""
    try:
        # Get default font from document styles
        default_style = doc.styles['Normal']
        default_font = {
            "family": default_style.font.name,
            "size": default_style.font.size.pt if default_style.font.size else None
        }
        
        # Get heading fonts
        heading_fonts = {}
        for i in range(1, 7):  # Heading 1 to Heading 6
            style_name = f'Heading {i}'
            if style_name in doc.styles:
                style = doc.styles[style_name]
                heading_fonts[f"h{i}"] = {
                    "family": style.font.name,
                    "size": style.font.size.pt if style.font.size else None
                }
        
        # Try to find code font
        code_font = None
        if 'Code' in doc.styles:
            style = doc.styles['Code']
            code_font = {
                "family": style.font.name,
                "size": style.font.size.pt if style.font.size else None
            }
        
        return {
            "body_font": default_font,
            "heading_fonts": heading_fonts,
            "code_font": code_font
        }
    
    except Exception as e:
        logger.error(f"Error extracting font settings: {str(e)}")
        return {"error": f"Error extracting font settings: {str(e)}"}

def check_toc_presence(doc):
    """Check for the presence of a table of contents"""
    try:
        has_toc = False
        toc_title = None
        
        # Look for TOC field codes
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if 'TOC' in run.text or 'Table of Contents' in run.text:
                    has_toc = True
                    # Try to find TOC title
                    if not toc_title:
                        # Check if this paragraph or the previous one might be the TOC title
                        if paragraph.style.name.startswith('Heading'):
                            toc_title = paragraph.text.strip()
                        elif paragraph.previous_paragraph and paragraph.previous_paragraph.style.name.startswith('Heading'):
                            toc_title = paragraph.previous_paragraph.text.strip()
        
        # If we found a TOC, try to determine its depth
        toc_depth = 0
        if has_toc:
            # Count the number of heading levels used in the document
            used_heading_levels = set()
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading '):
                    level = int(paragraph.style.name.split(' ')[1])
                    used_heading_levels.add(level)
            
            if used_heading_levels:
                toc_depth = max(used_heading_levels)
        
        return {
            "has_toc": has_toc,
            "toc_title": toc_title,
            "toc_depth": toc_depth
        }
    
    except Exception as e:
        logger.error(f"Error checking TOC presence: {str(e)}")
        return {"error": f"Error checking TOC presence: {str(e)}"}

def check_heading_numbering(doc):
    """Check heading numbering style (technical vs. standard)"""
    try:
        technical_count = 0
        standard_count = 0
        
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading '):
                text = paragraph.text.strip()
                
                # Check for technical numbering (1.1, 1.2, etc.)
                if re.match(r'^\d+(\.\d+)*\s+\w+', text):
                    technical_count += 1
                else:
                    standard_count += 1
        
        # Determine numbering style
        if technical_count > 0 and standard_count == 0:
            numbering_style = "technical"
        elif standard_count > 0 and technical_count == 0:
            numbering_style = "standard"
        elif technical_count > 0 and standard_count > 0:
            numbering_style = "mixed"
        else:
            numbering_style = "unknown"
        
        return {
            "heading_numbering": numbering_style,
            "technical_headings": technical_count,
            "standard_headings": standard_count
        }
    
    except Exception as e:
        logger.error(f"Error checking heading numbering: {str(e)}")
        return {"error": f"Error checking heading numbering: {str(e)}"}

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python docx_analyzer.py <docx_file>")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    result = analyze_docx(docx_file)
    
    import json
    print(json.dumps(result, indent=2))
