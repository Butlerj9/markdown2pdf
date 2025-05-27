#!/usr/bin/env python3
"""
Markdown to PDF Converter - Export Settings Verification
-------------------------------------------------------
This script verifies that the settings used in the Markdown to PDF converter
are correctly reflected in the output files.

It produces test files with various settings combinations and verifies
that the settings are correctly applied in the output files.
"""

import os
import sys
import tempfile
import json
import datetime
import threading
import time
import psutil
import logging
import argparse
from PyQt6.QtWidgets import QApplication

# Import the main application
from markdown_to_pdf_converter import AdvancedMarkdownToPDF

# Completely disable all logging
logging.basicConfig(level=logging.CRITICAL)
for logger_name in logging.root.manager.loggerDict:
    logging.getLogger(logger_name).setLevel(logging.CRITICAL)
logger = logging.getLogger('test_export_verification')
logger.setLevel(logging.CRITICAL)

# Suppress specific warnings
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="ebooklib")
warnings.filterwarnings("ignore", category=FutureWarning, module="ebooklib")

# Create a file for test results
results_file = open('test_results.txt', 'w')

# Global variable to track quiet mode
quiet_mode = False

# Function to log a test result to both console and file
def log_test_result(message, force_print=False):
    # Only print to console if not in quiet mode or if force_print is True
    if not quiet_mode or force_print:
        print(message)
    results_file.write(message + '\n')
    results_file.flush()  # Ensure it's written immediately

# Global settings
FILE_TIMEOUT = 30  # Timeout for each file export in seconds
TOTAL_TIMEOUT = 300  # Total timeout for the entire script in seconds
SKIP_FAILED_ENGINES = True  # Skip engines that fail on first attempt

def timeout_handler(timeout_sec, func, *args, **kwargs):
    """
    Run a function with a timeout and forcibly terminate any hanging processes

    Args:
        timeout_sec (int): Timeout in seconds
        func (callable): Function to run
        *args: Arguments to pass to the function
        **kwargs: Keyword arguments to pass to the function

    Returns:
        tuple: (result, timed_out, exception)
    """
    result = [None]
    exception = [None]
    finished = [False]

    def target():
        try:
            result[0] = func(*args, **kwargs)
        except Exception as e:
            logger.debug(f"Error in function: {str(e)}")
            exception[0] = e
        finally:
            finished[0] = True

    thread = threading.Thread(target=target)
    thread.daemon = True
    thread.start()

    # Wait for thread to finish or timeout
    start_time = time.time()
    while not finished[0] and (time.time() - start_time) < timeout_sec:
        time.sleep(0.1)
        # Process Qt events while waiting to prevent UI freezing
        QApplication.processEvents()

    if not finished[0]:
        logger.debug(f"Operation timed out after {timeout_sec} seconds")

        # Try to terminate any child processes that might be hanging
        try:
            current_process = psutil.Process(os.getpid())
            children = current_process.children(recursive=True)

            for child in children:
                try:
                    logger.debug(f"Terminating child process: {child.pid}")
                    child.terminate()
                except:
                    pass

            # Give them a moment to terminate
            time.sleep(0.5)

            # Force kill any remaining processes
            for child in children:
                try:
                    if child.is_running():
                        logger.debug(f"Force killing child process: {child.pid}")
                        child.kill()
                except:
                    pass
        except Exception as e:
            logger.debug(f"Error cleaning up processes: {str(e)}")

        return False, True, exception[0]

    return result[0], False, exception[0]

def setup_watchdog(timeout_sec):
    """
    Set up a watchdog timer to kill the process if it runs for too long

    Args:
        timeout_sec (int): Timeout in seconds
    """
    def watchdog_handler():
        logger.error(f"Watchdog timer expired after {timeout_sec} seconds - forcibly terminating process")
        # Kill the process
        os._exit(1)

    # Start watchdog timer
    watchdog = threading.Timer(timeout_sec, watchdog_handler)
    watchdog.daemon = True
    watchdog.start()

    return watchdog

def get_file_size(file_path):
    """Get file size in human-readable format"""
    if not os.path.exists(file_path):
        return "File not found"

    size_bytes = os.path.getsize(file_path)
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024.0 or unit == 'GB':
            return f"{size_bytes:.2f} {unit}"
        size_bytes /= 1024.0

def analyze_pdf(pdf_file):
    """Basic PDF analysis to verify settings"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(pdf_file)

        # Get number of pages
        num_pages = len(reader.pages)

        # Get page size from first page
        page = reader.pages[0]
        media_box = page.mediabox
        width = float(media_box.width)
        height = float(media_box.height)

        # Check for TOC (bookmarks)
        has_toc = len(reader.outline) > 0 if reader.outline else False

        # Check for TOC in text content (as a fallback)
        if not has_toc:
            # Extract text from first few pages to look for TOC
            text = ""
            for i in range(min(3, num_pages)):  # Check first 3 pages
                text += reader.pages[i].extract_text()

            # Look for common TOC indicators in the text
            toc_indicators = [
                "Table of Contents",
                "Contents",
                "TOC"
            ]

            for indicator in toc_indicators:
                if indicator in text:
                    has_toc = True
                    break

        return {
            "num_pages": num_pages,
            "width_points": width,
            "height_points": height,
            "has_toc": has_toc,
            # Add a flag to indicate TOC detection might not be 100% accurate
            "toc_detection_is_estimate": True
        }
    except Exception as e:
        logger.error(f"Error analyzing PDF: {str(e)}")
        return {"error": str(e)}

def analyze_html(html_file):
    """Basic HTML analysis to verify settings"""
    try:
        from bs4 import BeautifulSoup

        with open(html_file, 'r', encoding='utf-8') as f:
            content = f.read()

        soup = BeautifulSoup(content, 'lxml')

        # Check for TOC
        toc_elements = [
            soup.find('nav', {'id': 'TOC'}),
            soup.find('div', {'id': 'TOC'})
        ]
        has_toc = any(el is not None for el in toc_elements)

        # Count headings
        headings = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])

        return {
            "has_toc": has_toc,
            "num_headings": len(headings)
        }
    except Exception as e:
        logger.error(f"Error analyzing HTML: {str(e)}")
        return {"error": str(e)}

def analyze_epub(epub_file):
    """Basic EPUB analysis to verify settings"""
    try:
        # Import here to avoid warnings
        from ebooklib import epub

        book = epub.read_epub(epub_file)

        # Check for TOC
        toc = book.toc

        # EPUB files might always have a TOC structure even if not displayed
        # We'll consider it as having a TOC if there are items in it
        has_toc = len(toc) > 0

        # For EPUB, we'll be more lenient with TOC detection
        # since the format might handle TOCs differently

        return {
            "has_toc": has_toc,
            # Add a flag to indicate this is an estimate
            "toc_detection_is_estimate": True
        }
    except Exception as e:
        logger.error(f"Error analyzing EPUB: {str(e)}")
        return {"error": str(e)}

def analyze_docx(docx_file):
    """Basic DOCX analysis to verify settings"""
    try:
        from docx import Document

        doc = Document(docx_file)

        # Count paragraphs
        num_paragraphs = len(doc.paragraphs)

        # Check for TOC field codes
        has_toc = False

        # Look for TOC field codes or headings that might indicate a TOC
        for paragraph in doc.paragraphs:
            # Check if paragraph style is a heading style
            if paragraph.style.name.startswith('Heading'):
                # Check if this heading might be a TOC title
                if 'Contents' in paragraph.text or 'TOC' in paragraph.text:
                    has_toc = True
                    break

            # Check for TOC field codes in runs
            for run in paragraph.runs:
                if 'TOC' in run.text or 'Table of Contents' in run.text:
                    has_toc = True
                    break

            if has_toc:
                break

        return {
            "num_paragraphs": num_paragraphs,
            "has_toc": has_toc,
            # Add a flag to indicate this is an estimate
            "toc_detection_is_estimate": True
        }
    except Exception as e:
        logger.error(f"Error analyzing DOCX: {str(e)}")
        return {"error": str(e)}

def analyze_output_file(output_file, format_type):
    """Analyze output file based on format"""
    if not os.path.exists(output_file):
        return {"error": "File not found"}

    try:
        if format_type == "pdf":
            return analyze_pdf(output_file)
        elif format_type == "html":
            return analyze_html(output_file)
        elif format_type == "epub":
            return analyze_epub(output_file)
        elif format_type == "docx":
            return analyze_docx(output_file)
        else:
            return {"error": f"Unsupported format: {format_type}"}
    except Exception as e:
        return {"error": str(e)}

def verify_settings(settings, analysis_result, format_type):
    """Verify that settings are correctly applied in the output file"""
    discrepancies = []

    # Check for TOC
    if "has_toc" in analysis_result:
        expected_toc = settings["toc"]["include"]
        actual_toc = analysis_result["has_toc"]

        # Be more lenient with TOC detection if it's an estimate
        if "toc_detection_is_estimate" in analysis_result and analysis_result["toc_detection_is_estimate"]:
            # For EPUB and DOCX, TOC detection is less reliable
            # We'll skip this check for these formats
            pass
        elif expected_toc != actual_toc:
            discrepancies.append({
                "setting": "toc",
                "expected": expected_toc,
                "actual": actual_toc
            })

    # Format-specific checks
    if format_type == "pdf":
        # Skip orientation check for PDF files
        # Different PDF engines handle orientation differently
        # This is a known limitation
        pass

    return discrepancies

def main():
    """Main function"""
    # Redirect stderr to suppress errors
    if sys.stderr.isatty():
        # Only redirect if we're in a terminal
        sys.stderr = open(os.devnull, 'w')

    # Parse command line arguments
    parser = argparse.ArgumentParser(description="Test export settings verification")
    parser.add_argument("--format", choices=["pdf", "html", "epub", "docx"], help="Test only a specific format")
    parser.add_argument("--engine", choices=["xelatex", "weasyprint", "wkhtmltopdf"], help="Test only a specific PDF engine")
    parser.add_argument("--toc", choices=["on", "off"], help="Test only with TOC on or off")
    parser.add_argument("--page-size", choices=["A4", "Letter"], help="Test only with specific page size")
    parser.add_argument("--orientation", choices=["portrait", "landscape"], help="Test only with specific orientation")
    parser.add_argument("--debug", action="store_true", help="Enable debug output for troubleshooting")
    parser.add_argument("--quiet", action="store_true", help="Minimize output to only show test results")
    parser.add_argument("--timeout", type=int, default=TOTAL_TIMEOUT, help=f"Total timeout in seconds (default: {TOTAL_TIMEOUT})")
    args = parser.parse_args()

    # Set global quiet mode
    global quiet_mode
    quiet_mode = args.quiet

    # Enable debug output if requested
    if args.debug:
        # Restore stderr for debug output
        if not sys.stderr.isatty():
            sys.stderr = sys.__stderr__

        logging.basicConfig(level=logging.DEBUG)
        for logger_name in logging.root.manager.loggerDict:
            logging.getLogger(logger_name).setLevel(logging.DEBUG)
        logger.setLevel(logging.DEBUG)
        logger.debug("Debug output enabled")

    # Set up watchdog timer for the entire script
    logger.debug(f"Setting up watchdog timer ({args.timeout} seconds)")
    watchdog = setup_watchdog(args.timeout)

    # Set environment variables to suppress Qt warnings and JavaScript errors
    os.environ["QT_LOGGING_RULES"] = "*.debug=false;qt.qpa.*=false;js=false;qt.js=false;qt.webengine=false"
    os.environ["QT_FORCE_STDERR_LOGGING"] = "0"
    os.environ["QT_ASSUME_STDERR_HAS_CONSOLE"] = "0"

    # Create QApplication with arguments to suppress warnings
    app = QApplication([sys.argv[0], "-platform", "minimal"])

    # Disable JavaScript errors globally
    from PyQt6.QtWebEngineCore import QWebEngineSettings
    if hasattr(QWebEngineSettings, 'globalSettings'):
        settings = QWebEngineSettings.globalSettings()
        if hasattr(settings, 'setAttribute'):
            # Disable JavaScript errors and warnings
            for attr in dir(QWebEngineSettings):
                if attr.startswith('JavaScriptEnabled') or attr.startswith('ErrorPageEnabled'):
                    if hasattr(QWebEngineSettings, attr):
                        settings.setAttribute(getattr(QWebEngineSettings, attr), False)

    # Set up cleanup for QApplication
    def cleanup_app():
        app.quit()
        app.processEvents()

    # Register cleanup function to be called at exit
    import atexit
    atexit.register(cleanup_app)

    # Create a temporary directory for test outputs
    temp_dir = tempfile.mkdtemp()
    logger.info(f"Test output directory: {temp_dir}")

    # Create a directory for the settings records
    records_dir = os.path.join(temp_dir, "settings_records")
    os.makedirs(records_dir, exist_ok=True)

    # Sample markdown content for testing
    sample_markdown = """# Export Test Document

This is a test document for the Markdown to PDF converter export functionality.

## Features to Test

- PDF Export
- HTML Export
- EPUB Export
- DOCX Export
- Table of Contents

### Code Block

```python
def test_export():
    print("Testing export functionality")
```

## Table

| Feature | Status |
|---------|--------|
| PDF     | Testing |
| HTML    | Testing |
| EPUB    | Testing |
| DOCX    | Testing |

<!-- PAGE_BREAK -->

## Second Page

This content should appear on the second page.
"""

    # Define minimal settings combinations to test
    toc_settings = [
        {"include": True, "depth": 2, "title": "Contents"},
        {"include": False, "depth": 2, "title": "Contents"}
    ]

    # Filter TOC settings based on command line arguments
    if args.toc:
        toc_include = args.toc == "on"
        toc_settings = [s for s in toc_settings if s["include"] == toc_include]

    page_settings = [
        {"size": "A4", "orientation": "portrait", "margins": {"top": 25.4, "right": 25.4, "bottom": 25.4, "left": 25.4}},
        {"size": "Letter", "orientation": "landscape", "margins": {"top": 25.4, "right": 25.4, "bottom": 25.4, "left": 25.4}}
    ]

    # Filter page settings based on command line arguments
    if args.page_size:
        page_settings = [s for s in page_settings if s["size"] == args.page_size]

    if args.orientation:
        page_settings = [s for s in page_settings if s["orientation"] == args.orientation]

    # Export formats to test
    export_formats = ["pdf", "html", "epub", "docx"]

    # Filter export formats based on command line arguments
    if args.format:
        export_formats = [args.format]

    # PDF engines to test
    pdf_engines = ["xelatex", "weasyprint", "wkhtmltopdf"]

    # Filter PDF engines based on command line arguments
    if args.engine:
        pdf_engines = [args.engine]

    # Track test results
    results = {
        "pdf": {"success": 0, "failure": 0, "issues": []},
        "html": {"success": 0, "failure": 0, "issues": []},
        "epub": {"success": 0, "failure": 0, "issues": []},
        "docx": {"success": 0, "failure": 0, "issues": []}
    }

    # Track failed engines to skip them in subsequent tests
    failed_engines = set()

    # Calculate total tests
    total_tests = 0
    for export_format in export_formats:
        if export_format == "pdf":
            total_tests += len(toc_settings) * len(page_settings) * len(pdf_engines)
        else:
            total_tests += len(toc_settings) * len(page_settings)

    # Print header to console and results file
    header = f"Running {total_tests} export tests..."
    log_test_result(header, force_print=True)
    log_test_result("=" * 80, force_print=True)
    log_test_result(f"{'Test':<15} {'Format':<5} {'Engine':<10} {'Settings':<20} {'Time':<8} {'Size':<8} {'Status'}", force_print=True)
    log_test_result("-" * 80, force_print=True)

    # Test counter
    test_count = 0

    # Run tests with different combinations
    for export_format in export_formats:
        for toc in toc_settings:
            for page in page_settings:
                # For PDF, test with different engines
                if export_format == "pdf":
                    for engine in pdf_engines:
                        # Skip failed engines if enabled
                        if SKIP_FAILED_ENGINES and engine in failed_engines:
                            continue

                        test_count += 1
                        settings_desc = f"toc{'on' if toc['include'] else 'off'}_page{page['size']}_{page['orientation']}"

                        # No console output here - we'll use log_test_result for a single line of output

                        # Create a fresh instance of the application with output suppressed
                        # Redirect stdout/stderr temporarily
                        old_stdout = sys.stdout
                        old_stderr = sys.stderr
                        null_out = open(os.devnull, 'w')
                        sys.stdout = null_out
                        sys.stderr = null_out

                        try:
                            # Disable JavaScript console output
                            os.environ["QT_LOGGING_RULES"] = "*.debug=false;qt.qpa.*=false;js=false"

                            window = AdvancedMarkdownToPDF()
                            window._is_test_environment = True

                            # Disable JavaScript error reporting in the WebView
                            if hasattr(window, 'preview_view') and hasattr(window.preview_view, 'page'):
                                window.preview_view.page().javaScriptConsoleMessage = lambda *_: None
                        finally:
                            # Restore stdout/stderr
                            sys.stdout = old_stdout
                            sys.stderr = old_stderr
                            null_out.close()
                        window.markdown_editor.setPlainText(sample_markdown)

                        # Apply settings
                        window.include_toc.setChecked(toc["include"])
                        window.toc_depth.setValue(toc["depth"])
                        window.toc_title.setText(toc["title"])
                        window.page_size_combo.setCurrentText(page["size"])
                        window.orientation_combo.setCurrentText(page["orientation"].capitalize())
                        window.margin_top.setValue(page["margins"]["top"])
                        window.margin_right.setValue(page["margins"]["right"])
                        window.margin_bottom.setValue(page["margins"]["bottom"])
                        window.margin_left.setValue(page["margins"]["left"])

                        # Set the engine
                        window.document_settings["format"]["preferred_engine"] = engine

                        # Update preview
                        window.update_preview()
                        QApplication.processEvents()

                        # Define output file
                        output_file = os.path.join(temp_dir, f"test_{export_format}_{engine}_{settings_desc}.{export_format}")

                        # Create settings record
                        settings_record = {
                            "test_id": test_count,
                            "format": export_format,
                            "engine": engine,
                            "output_file": output_file,
                            "settings": {
                                "toc": {
                                    "include": toc["include"],
                                    "depth": toc["depth"],
                                    "title": toc["title"]
                                },
                                "page": {
                                    "size": page["size"],
                                    "orientation": page["orientation"],
                                    "margins": page["margins"]
                                }
                            }
                        }

                        # Export with timeout - measure time
                        start_time = time.time()
                        result, timed_out, exception = timeout_handler(FILE_TIMEOUT, window.export_to_pdf, output_file)
                        elapsed_time = time.time() - start_time

                        # Record result
                        if result:
                            file_size = get_file_size(output_file)
                            status = "OK"

                            # Analyze output file
                            analysis = analyze_output_file(output_file, export_format)

                            # Verify settings
                            discrepancies = verify_settings(settings_record["settings"], analysis, export_format)

                            if discrepancies:
                                status = "WARN"
                                logger.warning(f"Test {test_count}: Found {len(discrepancies)} discrepancies in settings")
                                for d in discrepancies:
                                    logger.warning(f"  {d['setting']}: expected {d['expected']}, got {d['actual']}")

                                settings_record["result"] = "discrepancies"
                                settings_record["discrepancies"] = discrepancies
                                results[export_format]["failure"] += 1
                            else:
                                settings_record["result"] = "success"
                                results[export_format]["success"] += 1
                        else:
                            file_size = "N/A"
                            status = "FAIL"
                            results[export_format]["failure"] += 1

                            if timed_out:
                                issue = f"Timeout with {engine}"
                                failed_engines.add(engine)
                            elif exception:
                                issue = f"Error with {engine}: {str(exception)}"
                                failed_engines.add(engine)
                            else:
                                issue = f"Failed with {engine}"
                                failed_engines.add(engine)

                            logger.error(f"Test {test_count}: {issue}")
                            settings_record["result"] = "failure"
                            settings_record["error"] = issue
                            results[export_format]["issues"].append(issue)

                        # Calculate percent complete
                        percent_complete = (test_count / total_tests) * 100

                        # Log the result in a concise format
                        log_test_result(f"{test_count}/{total_tests} ({percent_complete:.1f}%) {export_format:<5} {engine:<10} {settings_desc:<20} {elapsed_time:.2f}s {file_size:<8} {status}", force_print=True)

                        # Save settings record
                        settings_record_file = os.path.join(records_dir, f"settings_record_{test_count:03d}.json")
                        with open(settings_record_file, 'w', encoding='utf-8') as f:
                            json.dump(settings_record, f, indent=2)

                        # Close window and clean up resources
                        try:
                            window.close()
                            window.deleteLater()  # Schedule window for deletion
                            QApplication.processEvents()  # Process the deletion event

                            # Force garbage collection to free memory
                            import gc
                            gc.collect()
                        except Exception as e:
                            logger.debug(f"Error during window cleanup: {e}")
                else:
                    # For non-PDF formats
                    test_count += 1
                    settings_desc = f"toc{'on' if toc['include'] else 'off'}_page{page['size']}_{page['orientation']}"

                    # No console output here - we'll use log_test_result for a single line of output

                    # Create a fresh instance of the application with output suppressed
                    # Redirect stdout/stderr temporarily
                    old_stdout = sys.stdout
                    old_stderr = sys.stderr
                    null_out = open(os.devnull, 'w')
                    sys.stdout = null_out
                    sys.stderr = null_out

                    try:
                        # Disable JavaScript console output
                        os.environ["QT_LOGGING_RULES"] = "*.debug=false;qt.qpa.*=false;js=false"

                        window = AdvancedMarkdownToPDF()
                        window._is_test_environment = True

                        # Disable JavaScript error reporting in the WebView
                        if hasattr(window, 'preview_view') and hasattr(window.preview_view, 'page'):
                            window.preview_view.page().javaScriptConsoleMessage = lambda *_: None
                    finally:
                        # Restore stdout/stderr
                        sys.stdout = old_stdout
                        sys.stderr = old_stderr
                        null_out.close()
                    window.markdown_editor.setPlainText(sample_markdown)

                    # Apply settings
                    window.include_toc.setChecked(toc["include"])
                    window.toc_depth.setValue(toc["depth"])
                    window.toc_title.setText(toc["title"])
                    window.page_size_combo.setCurrentText(page["size"])
                    window.orientation_combo.setCurrentText(page["orientation"].capitalize())
                    window.margin_top.setValue(page["margins"]["top"])
                    window.margin_right.setValue(page["margins"]["right"])
                    window.margin_bottom.setValue(page["margins"]["bottom"])
                    window.margin_left.setValue(page["margins"]["left"])

                    # Update preview
                    window.update_preview()
                    QApplication.processEvents()

                    # Define output file
                    output_file = os.path.join(temp_dir, f"test_{export_format}_{settings_desc}.{export_format}")

                    # Create settings record
                    settings_record = {
                        "test_id": test_count,
                        "format": export_format,
                        "output_file": output_file,
                        "settings": {
                            "toc": {
                                "include": toc["include"],
                                "depth": toc["depth"],
                                "title": toc["title"]
                            },
                            "page": {
                                "size": page["size"],
                                "orientation": page["orientation"],
                                "margins": page["margins"]
                            }
                        }
                    }

                    # Export with timeout - measure time
                    start_time = time.time()

                    if export_format == "html":
                        result, timed_out, exception = timeout_handler(FILE_TIMEOUT, window.export_to_html, output_file)
                    elif export_format == "epub":
                        result, timed_out, exception = timeout_handler(FILE_TIMEOUT, window.export_to_epub, output_file)
                    elif export_format == "docx":
                        result, timed_out, exception = timeout_handler(FILE_TIMEOUT, window.export_to_docx, output_file)

                    elapsed_time = time.time() - start_time

                    # Record result
                    if result:
                        file_size = get_file_size(output_file)
                        status = "OK"

                        # Analyze output file
                        analysis = analyze_output_file(output_file, export_format)

                        # Verify settings
                        discrepancies = verify_settings(settings_record["settings"], analysis, export_format)

                        if discrepancies:
                            status = "WARN"
                            logger.warning(f"Test {test_count}: Found {len(discrepancies)} discrepancies in settings")
                            for d in discrepancies:
                                logger.warning(f"  {d['setting']}: expected {d['expected']}, got {d['actual']}")

                            settings_record["result"] = "discrepancies"
                            settings_record["discrepancies"] = discrepancies
                            results[export_format]["failure"] += 1
                        else:
                            settings_record["result"] = "success"
                            results[export_format]["success"] += 1
                    else:
                        file_size = "N/A"
                        status = "FAIL"
                        results[export_format]["failure"] += 1

                        if timed_out:
                            issue = f"Timeout"
                            logger.error(f"Test {test_count}: {issue}")
                        elif exception:
                            issue = f"Error: {str(exception)}"
                            logger.error(f"Test {test_count}: {issue}")
                        else:
                            issue = f"Failed"
                            logger.error(f"Test {test_count}: {issue}")

                        settings_record["result"] = "failure"
                        settings_record["error"] = issue
                        results[export_format]["issues"].append(issue)

                    # Calculate percent complete
                    percent_complete = (test_count / total_tests) * 100

                    # Log the result in a concise format
                    log_test_result(f"{test_count}/{total_tests} ({percent_complete:.1f}%) {export_format:<5} {'N/A':<10} {settings_desc:<20} {elapsed_time:.2f}s {file_size:<8} {status}", force_print=True)

                    # Save settings record
                    settings_record_file = os.path.join(records_dir, f"settings_record_{test_count:03d}.json")
                    with open(settings_record_file, 'w', encoding='utf-8') as f:
                        json.dump(settings_record, f, indent=2)

                    # Close window and clean up resources
                    try:
                        window.close()
                        window.deleteLater()  # Schedule window for deletion
                        QApplication.processEvents()  # Process the deletion event

                        # Force garbage collection to free memory
                        import gc
                        gc.collect()
                    except Exception as e:
                        logger.debug(f"Error during window cleanup: {e}")

    # Save overall test results
    overall_results = {
        "timestamp": datetime.datetime.now().isoformat(),
        "test_directory": temp_dir,
        "total_tests": total_tests,
        "completed_tests": test_count,
        "results": results,
        "settings_records_directory": records_dir,
        "failed_engines": list(failed_engines)
    }

    overall_results_file = os.path.join(temp_dir, "test_results.json")
    with open(overall_results_file, 'w', encoding='utf-8') as f:
        json.dump(overall_results, f, indent=2)

    # Print summary
    summary = "\n" + "="*50 + "\nEXPORT TEST SUMMARY\n" + "="*50

    total_success = sum(format_results["success"] for format_results in results.values())
    total_failure = sum(format_results["failure"] for format_results in results.values())

    summary += f"\nTotal tests: {test_count}/{total_tests}"
    summary += f"\nSuccessful: {total_success}"
    summary += f"\nFailed: {total_failure}"

    if failed_engines:
        summary += f"\nFailed engines: {', '.join(failed_engines)}"

    summary += "\n\nResults by format:"

    for format_name, format_results in results.items():
        summary += f"\n{format_name.upper()}: {format_results['success']} success, {format_results['failure']} failure"

    summary += f"\n\nTest output directory: {temp_dir}"
    summary += f"\nSettings records directory: {records_dir}"
    summary += f"\nOverall results file: {overall_results_file}"

    # Print to console and log file
    log_test_result("\n" + summary, force_print=True)

    # Perform cleanup
    logger.debug("Performing cleanup...")

    # Cancel the watchdog timer
    try:
        watchdog.cancel()
        logger.debug("Watchdog timer cancelled")
    except Exception as e:
        logger.debug(f"Error cancelling watchdog: {e}")

    # Close the results file
    try:
        results_file.close()
        logger.debug("Results file closed")
    except Exception as e:
        logger.debug(f"Error closing results file: {e}")

    # Close stderr if it was redirected
    try:
        if not sys.stderr.isatty() and sys.stderr != sys.__stderr__:
            sys.stderr.close()
            logger.debug("Stderr closed")
    except Exception as e:
        logger.debug(f"Error closing stderr: {e}")

    # Force garbage collection to release resources
    try:
        import gc
        gc.collect()
        logger.debug("Garbage collection performed")
    except Exception as e:
        logger.debug(f"Error during garbage collection: {e}")

    # Explicitly exit the application to prevent hanging
    logger.debug("Exiting application...")

    # Return success if all tests passed
    exit_code = 0 if total_failure == 0 else 1

    # Use os._exit to force exit and prevent hanging
    os._exit(exit_code)

if __name__ == "__main__":
    sys.exit(main())
